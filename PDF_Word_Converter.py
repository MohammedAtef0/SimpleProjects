# First, you need to run those bash commands:
# pip install PyPDF2 python-docx
# pip install PyMuPDF python-docx
# -------------------------------------------------------------------------------------------------------------------------
from PyPDF2 import PdfReader
from docx import Document

def pdf_to_word_basic(pdf_path, word_path):
    """
    Converts a PDF file to a Word document (.docx) by extracting text.
    Note: This is a basic conversion and will not preserve formatting,
    images, or complex layouts. It primarily extracts text.
    """
    try:
        # Create a new Word document
        document = Document()

        # Open the PDF file
        with open(pdf_path, 'rb') as file:
            reader = PdfReader(file)
            print(f"Number of pages in PDF: {len(reader.pages)}")

            # Iterate through each page and extract text
            for page_num in range(len(reader.pages)):
                page = reader.pages[page_num]
                text = page.extract_text()

                if text:
                    # Add extracted text to the Word document
                    document.add_paragraph(text)
                else:
                    document.add_paragraph(f"--- No text found on page {page_num + 1} ---")

        # Save the Word document
        document.save(word_path)
        print(f"Successfully converted '{pdf_path}' to '{word_path}'")

    except Exception as e:
        print(f"An error occurred: {e}")

# --- Example Usage ---
if __name__ == "__main__":
    # Create a dummy PDF file for testing (you'll need a real PDF)
    # For demonstration, let's assume you have a 'sample.pdf'
    # If you don't have one, you can create a simple text file and convert it to PDF online,
    # or just use an existing PDF.
    
    # You will need to replace 'input.pdf' with the path to your actual PDF file
    input_pdf_file = "input.pdf" 
    output_word_file = "output.docx"

    # Create a dummy input.pdf for testing purposes if it doesn't exist
    try:
        from reportlab.pdfgen import canvas
        c = canvas.Canvas(input_pdf_file)
        c.drawString(100, 750, "Hello, this is a sample PDF.")
        c.drawString(100, 730, "This is the second line of text.")
        c.showPage()
        c.drawString(100, 750, "This is another page.")
        c.save()
        print(f"Created a dummy PDF: {input_pdf_file}")
    except ImportError:
        print("reportlab not installed. Please ensure you have 'input.pdf' for testing.")
        print("You can install reportlab with: pip install reportlab")
    except Exception as e:
        print(f"Could not create dummy PDF (perhaps file already exists or permissions issue): {e}")

    pdf_to_word_basic(input_pdf_file, output_word_file)
# ================================================================================================================================================
